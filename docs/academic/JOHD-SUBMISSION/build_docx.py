"""Genera il manoscritto JOHD in .docx dal markdown.

JOHD accetta .doc(x), NON markdown né PDF (per submission Word).
Struttura: quella dei Data Paper JOHD (Abstract + sezioni 1-4 + dichiarazioni +
References in stile APA).

Uso:  python docs/academic/JOHD-SUBMISSION/build_docx.py
Output: docs/academic/JOHD-SUBMISSION/AtlasPI-JOHD-data-paper.docx
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

HERE = Path(__file__).resolve().parent
SRC = HERE / "manuscript.md"
OUT = HERE / "AtlasPI-JOHD-data-paper.docx"


def strip_md(text: str) -> str:
    """Rimuove la sintassi markdown inline, preservando il testo."""
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)  # link -> testo (url)
    text = text.replace("**", "").replace("`", "")
    text = re.sub(r"(?<!\w)\*(?!\s)(.+?)(?<!\s)\*(?!\w)", r"\1", text)  # *corsivo*
    return text.strip()


def main() -> None:
    md = SRC.read_text(encoding="utf-8")
    # scarta l'eventuale blocco di note iniziale (righe che iniziano con ">")
    lines = [ln for ln in md.splitlines() if not ln.lstrip().startswith(">")]

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    in_table = False
    table_rows: list[list[str]] = []

    def flush_table() -> None:
        """Scrive la tabella accumulata (i separatori |---| vengono scartati)."""
        nonlocal table_rows
        rows = [r for r in table_rows if not all(set(c) <= set("-: ") for c in r)]
        if rows:
            t = doc.add_table(rows=len(rows), cols=len(rows[0]))
            t.style = "Table Grid"
            for i, row in enumerate(rows):
                for j, cell in enumerate(row):
                    t.cell(i, j).text = strip_md(cell)
        table_rows = []

    for raw in lines:
        line = raw.rstrip()

        if line.startswith("|"):
            in_table = True
            table_rows.append([c.strip() for c in line.strip("|").split("|")])
            continue
        if in_table:
            flush_table()
            in_table = False

        if not line.strip() or line.strip() == "---":
            continue

        if line.startswith("# "):
            h = doc.add_heading(strip_md(line[2:]), level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(strip_md(line[3:]), level=1)
        elif line.startswith("### "):
            doc.add_heading(strip_md(line[4:]), level=2)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(strip_md(line[2:]), style="List Bullet")
        else:
            p = doc.add_paragraph(strip_md(line))
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if in_table:
        flush_table()

    doc.save(OUT)
    print(f"scritto: {OUT}")


if __name__ == "__main__":
    main()
